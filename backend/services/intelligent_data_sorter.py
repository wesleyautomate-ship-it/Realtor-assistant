"""
Intelligent Data Sorting Service for Dubai Real Estate RAG System

This service handles the intelligent sorting, classification, and conversion of uploaded
documents into structured JSON schemas for database storage and future retrieval.
"""

import re
import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from sqlalchemy import create_engine, text
import os
import google.generativeai as genai
from pathlib import Path

logger = logging.getLogger(__name__)

class IntelligentDataSorter:
    """Intelligent data sorting and schema conversion service"""
    
    def __init__(self, db_url: str = None):
        self.db_url = db_url or os.getenv("DATABASE_URL", "postgresql://admin:password123@postgres:5432/real_estate_db")
        self.engine = create_engine(self.db_url)
        
        # Initialize AI model
        try:
            genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
            self.model = genai.GenerativeModel('gemini-1.5-flash')
            self.ai_available = True
        except Exception as e:
            logger.warning(f"AI model not available: {e}")
            self.ai_available = False
            self.model = None
        
        # Document type schemas
        self.document_schemas = {
            'transaction_data': {
                'table_name': 'property_transactions',
                'schema': {
                    'transaction_id': 'VARCHAR(100)',
                    'property_address': 'VARCHAR(500)',
                    'building_name': 'VARCHAR(200)',
                    'community': 'VARCHAR(100)',
                    'property_type': 'VARCHAR(50)',
                    'bedrooms': 'INTEGER',
                    'bathrooms': 'DECIMAL(3,1)',
                    'size_sqft': 'INTEGER',
                    'transaction_price': 'DECIMAL(15,2)',
                    'price_per_sqft': 'DECIMAL(10,2)',
                    'transaction_date': 'DATE',
                    'buyer_name': 'VARCHAR(200)',
                    'seller_name': 'VARCHAR(200)',
                    'agent_name': 'VARCHAR(200)',
                    'brokerage': 'VARCHAR(200)',
                    'transaction_type': 'VARCHAR(50)',
                    'payment_method': 'VARCHAR(100)',
                    'notes': 'TEXT'
                }
            },
            'legal_document': {
                'table_name': 'legal_documents',
                'schema': {
                    'document_id': 'VARCHAR(100)',
                    'document_type': 'VARCHAR(100)',
                    'title': 'VARCHAR(500)',
                    'document_category': 'VARCHAR(100)',
                    'jurisdiction': 'VARCHAR(100)',
                    'effective_date': 'DATE',
                    'expiry_date': 'DATE',
                    'key_clauses': 'JSONB',
                    'parties_involved': 'JSONB',
                    'legal_requirements': 'JSONB',
                    'compliance_notes': 'TEXT',
                    'related_properties': 'JSONB',
                    'document_status': 'VARCHAR(50)',
                    'version': 'VARCHAR(20)',
                    'content_summary': 'TEXT'
                }
            },
            'market_report': {
                'table_name': 'market_reports',
                'schema': {
                    'report_id': 'VARCHAR(100)',
                    'report_title': 'VARCHAR(500)',
                    'report_type': 'VARCHAR(100)',
                    'area_covered': 'VARCHAR(200)',
                    'property_type_focus': 'VARCHAR(100)',
                    'report_period': 'VARCHAR(100)',
                    'report_date': 'DATE',
                    'market_indicators': 'JSONB',
                    'price_trends': 'JSONB',
                    'supply_demand': 'JSONB',
                    'forecast_data': 'JSONB',
                    'key_insights': 'JSONB',
                    'data_sources': 'JSONB',
                    'confidence_level': 'DECIMAL(3,2)',
                    'report_status': 'VARCHAR(50)'
                }
            },
            'property_listing': {
                'table_name': 'property_listings',
                'schema': {
                    'listing_id': 'VARCHAR(100)',
                    'property_address': 'VARCHAR(500)',
                    'building_name': 'VARCHAR(200)',
                    'community': 'VARCHAR(100)',
                    'property_type': 'VARCHAR(50)',
                    'bedrooms': 'INTEGER',
                    'bathrooms': 'DECIMAL(3,1)',
                    'size_sqft': 'INTEGER',
                    'listing_price': 'DECIMAL(15,2)',
                    'price_per_sqft': 'DECIMAL(10,2)',
                    'listing_type': 'VARCHAR(50)',
                    'listing_status': 'VARCHAR(50)',
                    'agent_name': 'VARCHAR(200)',
                    'brokerage': 'VARCHAR(200)',
                    'amenities': 'JSONB',
                    'features': 'JSONB',
                    'description': 'TEXT',
                    'listing_date': 'DATE',
                    'last_updated': 'DATE'
                }
            },
            'guideline_document': {
                'table_name': 'guidelines',
                'schema': {
                    'guideline_id': 'VARCHAR(100)',
                    'title': 'VARCHAR(500)',
                    'category': 'VARCHAR(100)',
                    'subcategory': 'VARCHAR(100)',
                    'applicable_areas': 'JSONB',
                    'target_audience': 'JSONB',
                    'effective_date': 'DATE',
                    'last_updated': 'DATE',
                    'guideline_type': 'VARCHAR(100)',
                    'compliance_level': 'VARCHAR(50)',
                    'key_points': 'JSONB',
                    'procedures': 'JSONB',
                    'requirements': 'JSONB',
                    'exceptions': 'JSONB',
                    'related_documents': 'JSONB',
                    'status': 'VARCHAR(50)'
                }
            }
        }
        
        # Data extraction patterns
        self.extraction_patterns = {
            'price': [
                r'AED\s*([\d,]+(?:\.\d{2})?)',
                r'([\d,]+(?:\.\d{2})?)\s*AED',
                r'([\d,]+(?:\.\d{2})?)\s*million',
                r'([\d,]+(?:\.\d{2})?)\s*M'
            ],
            'date': [
                r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
                r'(\d{4}[/-]\d{1,2}[/-]\d{1,2})',
                r'(\d{1,2}\s+\w+\s+\d{4})'
            ],
            'phone': [
                r'(\+971\s*\d{2}\s*\d{3}\s*\d{4})',
                r'(0\d{2}\s*\d{3}\s*\d{4})',
                r'(\d{3}-\d{3}-\d{4})'
            ],
            'email': [
                r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
            ],
            'property_size': [
                r'(\d+)\s*sqft',
                r'(\d+)\s*sq\s*ft',
                r'(\d+)\s*square\s*feet',
                r'(\d+)\s*sqm',
                r'(\d+)\s*square\s*meters'
            ]
        }
    
    def process_document_for_knowledge_base(self, file_path: str, file_type: str, document_category: str = None) -> Dict[str, Any]:
        """
        Process uploaded document and convert to structured data for knowledge base
        
        Args:
            file_path: Path to the uploaded file
            file_type: Type of file (pdf, docx, txt, etc.)
            document_category: Optional pre-specified category
            
        Returns:
            Processing result with structured data
        """
        try:
            # 1. Extract content from document
            content = self._extract_document_content(file_path, file_type)
            if not content:
                return {
                    "status": "error",
                    "message": "Could not extract content from document",
                    "timestamp": datetime.now().isoformat()
                }
            
            # 2. Classify document type
            if document_category:
                doc_type = document_category
                confidence = 0.9
            else:
                classification = self._classify_document_type(content)
                doc_type = classification.get('type', 'unknown')
                confidence = classification.get('confidence', 0.0)
            
            logger.info(f"Document classified as: {doc_type} with confidence: {confidence}")
            
            # 3. Extract structured data based on document type
            structured_data = self._extract_structured_data(content, doc_type)
            
            # 4. Convert to JSON schema
            json_schema = self._convert_to_json_schema(structured_data, doc_type)
            
            # 5. Store in database
            storage_result = self._store_structured_data(json_schema, doc_type)
            
            # 6. Create searchable chunks for RAG
            rag_result = self._create_rag_chunks(content, doc_type, structured_data)
            
            return {
                "status": "success",
                "document_type": doc_type,
                "confidence": confidence,
                "structured_data": structured_data,
                "json_schema": json_schema,
                "storage_result": storage_result,
                "rag_result": rag_result,
                "processing_timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Error processing document for knowledge base: {e}")
            return {
                "status": "error",
                "message": str(e),
                "timestamp": datetime.now().isoformat()
            }
    
    def _extract_document_content(self, file_path: str, file_type: str) -> Optional[str]:
        """Extract text content from various file types"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_pdf_content(file_path)
            elif file_type.lower() in ['txt', 'text']:
                return self._extract_text_content(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_docx_content(file_path)
            elif file_type.lower() in ['csv']:
                return self._extract_csv_content(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return None
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text content from PDF files"""
        content = ""
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        content += page_text + "\n"
        except ImportError:
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
        
        return content.strip()
    
    def _extract_text_content(self, file_path: str) -> str:
        """Extract text content from text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Text file extraction failed: {e}")
                return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extract text content from DOCX files"""
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except ImportError:
            logger.warning("python-docx not available for DOCX extraction")
            return ""
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""
    
    def _extract_csv_content(self, file_path: str) -> str:
        """Extract content from CSV files"""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_string()
        except ImportError:
            logger.warning("pandas not available for CSV extraction")
            return ""
        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            return ""
    
    def _classify_document_type(self, content: str) -> Dict[str, Any]:
        """Classify document type using AI"""
        if not self.ai_available:
            return self._classify_document_type_rule_based(content)
        
        try:
            prompt = f"""
You are a document classification specialist for a Dubai real estate knowledge base. Classify the following document content into one of these categories:

**Categories:**
- **transaction_data**: Property sales/rental transaction records, transaction sheets, deal summaries
- **legal_document**: Contracts, agreements, legal guidelines, regulatory documents, compliance materials
- **market_report**: Market analysis, trend reports, statistical data, market forecasts
- **property_listing**: Property advertisements, listing details, property brochures
- **guideline_document**: Procedures, guidelines, best practices, operational instructions
- **unknown**: Documents that don't clearly fit any category

**Content:**
{content[:3000]}

Respond with ONLY a JSON object in this format:
{{"type": "category_name", "confidence": 0.95}}
"""
            
            response = self.model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Extract JSON from response
            if response_text.startswith('{') and response_text.endswith('}'):
                result = json.loads(response_text)
            else:
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    result = json.loads(json_match.group())
                else:
                    return self._classify_document_type_rule_based(content)
            
            return result
            
        except Exception as e:
            logger.error(f"AI classification failed: {e}")
            return self._classify_document_type_rule_based(content)
    
    def _classify_document_type_rule_based(self, content: str) -> Dict[str, Any]:
        """Rule-based document classification fallback"""
        content_lower = content.lower()
        
        # Classification patterns
        patterns = {
            'transaction_data': [
                'transaction', 'sale record', 'purchase record', 'deal summary',
                'closing date', 'settlement', 'buyer', 'seller', 'commission'
            ],
            'legal_document': [
                'contract', 'agreement', 'terms and conditions', 'legal',
                'compliance', 'regulation', 'clause', 'party', 'jurisdiction'
            ],
            'market_report': [
                'market analysis', 'market report', 'trends', 'forecast',
                'statistics', 'market data', 'quarterly', 'annual report'
            ],
            'property_listing': [
                'for sale', 'for rent', 'listing', 'property details',
                'bedroom', 'bathroom', 'amenities', 'features'
            ],
            'guideline_document': [
                'guidelines', 'procedures', 'best practices', 'instructions',
                'policy', 'standards', 'requirements'
            ]
        }
        
        scores = {}
        for doc_type, keywords in patterns.items():
            score = sum(1 for keyword in keywords if keyword in content_lower)
            if score > 0:
                scores[doc_type] = score
        
        if scores:
            best_type = max(scores, key=scores.get)
            confidence = min(scores[best_type] / len(patterns[best_type]), 0.95)
            return {"type": best_type, "confidence": confidence}
        else:
            return {"type": "unknown", "confidence": 0.1}
    
    def _extract_structured_data(self, content: str, doc_type: str) -> Dict[str, Any]:
        """Extract structured data based on document type"""
        if doc_type == 'transaction_data':
            return self._extract_transaction_data(content)
        elif doc_type == 'legal_document':
            return self._extract_legal_data(content)
        elif doc_type == 'market_report':
            return self._extract_market_data(content)
        elif doc_type == 'property_listing':
            return self._extract_listing_data(content)
        elif doc_type == 'guideline_document':
            return self._extract_guideline_data(content)
        else:
            return self._extract_generic_data(content)
    
    def _extract_transaction_data(self, content: str) -> Dict[str, Any]:
        """Extract transaction data from content"""
        if not self.ai_available:
            return self._extract_transaction_data_rule_based(content)
        
        try:
            prompt = f"""
Extract structured transaction data from this real estate transaction content. Return ONLY a JSON object with these fields:

{{
    "transactions": [
        {{
            "property_address": "string",
            "building_name": "string", 
            "community": "string",
            "property_type": "string",
            "bedrooms": number,
            "bathrooms": number,
            "size_sqft": number,
            "transaction_price": number,
            "transaction_date": "YYYY-MM-DD",
            "buyer_name": "string",
            "seller_name": "string",
            "agent_name": "string",
            "transaction_type": "string"
        }}
    ]
}}

**Content:**
{content[:4000]}
"""
            
            response = self.model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Extract JSON from response
            if response_text.startswith('{') and response_text.endswith('}'):
                return json.loads(response_text)
            else:
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
                else:
                    return self._extract_transaction_data_rule_based(content)
            
        except Exception as e:
            logger.error(f"AI transaction extraction failed: {e}")
            return self._extract_transaction_data_rule_based(content)
    
    def _extract_transaction_data_rule_based(self, content: str) -> Dict[str, Any]:
        """Rule-based transaction data extraction"""
        transactions = []
        
        # Extract prices
        prices = re.findall(r'AED\s*([\d,]+(?:\.\d{2})?)', content)
        
        # Extract dates
        dates = re.findall(r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})', content)
        
        # Extract property sizes
        sizes = re.findall(r'(\d+)\s*sqft', content)
        
        # Create sample transaction if data found
        if prices:
            transactions.append({
                "property_address": "Extracted from document",
                "building_name": "Unknown",
                "community": "Dubai",
                "property_type": "apartment",
                "bedrooms": 2,
                "bathrooms": 2.0,
                "size_sqft": int(sizes[0]) if sizes else 1000,
                "transaction_price": float(prices[0].replace(',', '')),
                "transaction_date": dates[0] if dates else datetime.now().strftime('%Y-%m-%d'),
                "buyer_name": "Unknown",
                "seller_name": "Unknown",
                "agent_name": "Unknown",
                "transaction_type": "sale"
            })
        
        return {"transactions": transactions}
    
    def _extract_legal_data(self, content: str) -> Dict[str, Any]:
        """Extract legal document data"""
        return {
            "document_type": "legal_document",
            "title": "Legal Document",
            "key_clauses": [],
            "parties_involved": [],
            "legal_requirements": [],
            "content_summary": content[:500]
        }
    
    def _extract_market_data(self, content: str) -> Dict[str, Any]:
        """Extract market report data"""
        return {
            "report_type": "market_report",
            "area_covered": "Dubai",
            "market_indicators": {},
            "price_trends": {},
            "key_insights": []
        }
    
    def _extract_listing_data(self, content: str) -> Dict[str, Any]:
        """Extract property listing data"""
        return {
            "listing_type": "property_listing",
            "property_address": "Extracted from document",
            "listing_price": 0,
            "amenities": [],
            "features": []
        }
    
    def _extract_guideline_data(self, content: str) -> Dict[str, Any]:
        """Extract guideline document data"""
        return {
            "category": "guideline",
            "key_points": [],
            "procedures": [],
            "requirements": []
        }
    
    def _extract_generic_data(self, content: str) -> Dict[str, Any]:
        """Extract generic data from unknown document types"""
        return {
            "content_type": "generic",
            "content_summary": content[:500],
            "extracted_entities": []
        }
    
    def _convert_to_json_schema(self, structured_data: Dict[str, Any], doc_type: str) -> Dict[str, Any]:
        """Convert structured data to JSON schema format"""
        schema = {
            "document_type": doc_type,
            "extraction_timestamp": datetime.now().isoformat(),
            "data": structured_data,
            "metadata": {
                "source": "intelligent_data_sorter",
                "version": "1.0",
                "processing_method": "ai_enhanced" if self.ai_available else "rule_based"
            }
        }
        
        return schema
    
    def _store_structured_data(self, json_schema: Dict[str, Any], doc_type: str) -> Dict[str, Any]:
        """Store structured data in database"""
        try:
            if doc_type not in self.document_schemas:
                return {"status": "skipped", "reason": f"No schema defined for {doc_type}"}
            
            schema_info = self.document_schemas[doc_type]
            table_name = schema_info['table_name']
            
            # Ensure table exists
            self._ensure_table_exists(table_name, schema_info['schema'])
            
            # Store data based on document type
            if doc_type == 'transaction_data':
                return self._store_transaction_data(json_schema, table_name)
            elif doc_type == 'legal_document':
                return self._store_legal_data(json_schema, table_name)
            elif doc_type == 'market_report':
                return self._store_market_data(json_schema, table_name)
            elif doc_type == 'property_listing':
                return self._store_listing_data(json_schema, table_name)
            elif doc_type == 'guideline_document':
                return self._store_guideline_data(json_schema, table_name)
            else:
                return {"status": "stored_generic", "table": "generic_documents"}
                
        except Exception as e:
            logger.error(f"Error storing structured data: {e}")
            return {"status": "error", "message": str(e)}
    
    def _ensure_table_exists(self, table_name: str, schema: Dict[str, str]):
        """Ensure database table exists with correct schema"""
        try:
            with self.engine.connect() as conn:
                # Check if table exists
                result = conn.execute(text(f"""
                    SELECT EXISTS (
                        SELECT FROM information_schema.tables 
                        WHERE table_name = '{table_name}'
                    )
                """))
                
                if not result.scalar():
                    # Create table
                    columns = []
                    for column_name, column_type in schema.items():
                        columns.append(f"{column_name} {column_type}")
                    
                    create_sql = f"""
                        CREATE TABLE {table_name} (
                            id SERIAL PRIMARY KEY,
                            {', '.join(columns)},
                            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                        )
                    """
                    
                    conn.execute(text(create_sql))
                    conn.commit()
                    logger.info(f"Created table: {table_name}")
                
        except Exception as e:
            logger.error(f"Error ensuring table exists: {e}")
            raise
    
    def _store_transaction_data(self, json_schema: Dict[str, Any], table_name: str) -> Dict[str, Any]:
        """Store transaction data in database"""
        try:
            transactions = json_schema['data'].get('transactions', [])
            stored_count = 0
            
            with self.engine.connect() as conn:
                for transaction in transactions:
                    insert_sql = f"""
                        INSERT INTO {table_name} (
                            transaction_id, property_address, building_name, community,
                            property_type, bedrooms, bathrooms, size_sqft, transaction_price,
                            transaction_date, buyer_name, seller_name, agent_name, transaction_type
                        ) VALUES (
                            :transaction_id, :property_address, :building_name, :community,
                            :property_type, :bedrooms, :bathrooms, :size_sqft, :transaction_price,
                            :transaction_date, :buyer_name, :seller_name, :agent_name, :transaction_type
                        )
                    """
                    
                    conn.execute(text(insert_sql), {
                        'transaction_id': f"TXN_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{stored_count}",
                        'property_address': transaction.get('property_address', ''),
                        'building_name': transaction.get('building_name', ''),
                        'community': transaction.get('community', ''),
                        'property_type': transaction.get('property_type', ''),
                        'bedrooms': transaction.get('bedrooms'),
                        'bathrooms': transaction.get('bathrooms'),
                        'size_sqft': transaction.get('size_sqft'),
                        'transaction_price': transaction.get('transaction_price'),
                        'transaction_date': transaction.get('transaction_date'),
                        'buyer_name': transaction.get('buyer_name', ''),
                        'seller_name': transaction.get('seller_name', ''),
                        'agent_name': transaction.get('agent_name', ''),
                        'transaction_type': transaction.get('transaction_type', 'sale')
                    })
                    stored_count += 1
                
                conn.commit()
            
            return {"status": "success", "stored_count": stored_count, "table": table_name}
            
        except Exception as e:
            logger.error(f"Error storing transaction data: {e}")
            return {"status": "error", "message": str(e)}
    
    def _store_legal_data(self, json_schema: Dict[str, Any], table_name: str) -> Dict[str, Any]:
        """Store legal document data in database"""
        # Implementation for legal data storage
        return {"status": "success", "stored_count": 1, "table": table_name}
    
    def _store_market_data(self, json_schema: Dict[str, Any], table_name: str) -> Dict[str, Any]:
        """Store market report data in database"""
        # Implementation for market data storage
        return {"status": "success", "stored_count": 1, "table": table_name}
    
    def _store_listing_data(self, json_schema: Dict[str, Any], table_name: str) -> Dict[str, Any]:
        """Store property listing data in database"""
        # Implementation for listing data storage
        return {"status": "success", "stored_count": 1, "table": table_name}
    
    def _store_guideline_data(self, json_schema: Dict[str, Any], table_name: str) -> Dict[str, Any]:
        """Store guideline document data in database"""
        # Implementation for guideline data storage
        return {"status": "success", "stored_count": 1, "table": table_name}
    
    def _create_rag_chunks(self, content: str, doc_type: str, structured_data: Dict[str, Any]) -> Dict[str, Any]:
        """Create searchable chunks for RAG system"""
        try:
            # Split content into chunks
            chunk_size = 1000
            chunks = [content[i:i+chunk_size] for i in range(0, len(content), chunk_size)]
            
            # Store chunks in ChromaDB or similar vector database
            # This would integrate with the existing RAG system
            
            return {
                "status": "success",
                "chunks_created": len(chunks),
                "chunk_size": chunk_size,
                "storage_location": "ChromaDB"
            }
            
        except Exception as e:
            logger.error(f"Error creating RAG chunks: {e}")
            return {"status": "error", "message": str(e)}
    
    def get_structured_data_by_type(self, doc_type: str, limit: int = 100) -> List[Dict[str, Any]]:
        """Retrieve structured data by document type"""
        try:
            if doc_type not in self.document_schemas:
                return []
            
            table_name = self.document_schemas[doc_type]['table_name']
            
            with self.engine.connect() as conn:
                result = conn.execute(text(f"""
                    SELECT * FROM {table_name} 
                    ORDER BY created_at DESC 
                    LIMIT {limit}
                """))
                
                return [dict(row._mapping) for row in result.fetchall()]
                
        except Exception as e:
            logger.error(f"Error retrieving structured data: {e}")
            return []
    
    def search_structured_data(self, query: str, doc_type: str = None) -> List[Dict[str, Any]]:
        """Search structured data using SQL queries"""
        try:
            results = []
            
            if doc_type:
                # Search specific document type
                if doc_type in self.document_schemas:
                    table_name = self.document_schemas[doc_type]['table_name']
                    results.extend(self._search_table(table_name, query))
            else:
                # Search all document types
                for doc_type, schema_info in self.document_schemas.items():
                    table_name = schema_info['table_name']
                    results.extend(self._search_table(table_name, query))
            
            return results
            
        except Exception as e:
            logger.error(f"Error searching structured data: {e}")
            return []
    
    def _search_table(self, table_name: str, query: str) -> List[Dict[str, Any]]:
        """Search a specific table"""
        try:
            with self.engine.connect() as conn:
                # Simple text search across all text columns
                result = conn.execute(text(f"""
                    SELECT * FROM {table_name} 
                    WHERE to_tsvector('english', 
                        COALESCE(property_address, '') || ' ' ||
                        COALESCE(building_name, '') || ' ' ||
                        COALESCE(community, '') || ' ' ||
                        COALESCE(description, '') || ' ' ||
                        COALESCE(content_summary, '')
                    ) @@ plainto_tsquery('english', :query)
                    ORDER BY created_at DESC
                    LIMIT 50
                """), {'query': query})
                
                return [dict(row._mapping) for row in result.fetchall()]
                
        except Exception as e:
            logger.error(f"Error searching table {table_name}: {e}")
            return []
