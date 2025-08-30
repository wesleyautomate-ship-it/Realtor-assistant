#!/usr/bin/env python3
"""
Sample Data Generator for Dubai Real Estate RAG Application
Generates comprehensive datasets for testing at scale
"""

import pandas as pd
import numpy as np
import random
import json
import csv
from datetime import datetime, timedelta
import os
from faker import Faker
import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
import docx
from docx import Document
from docx.shared import Inches
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io

# Initialize Faker for realistic data
fake = Faker(['en_US', 'ar_SA'])

# Dubai neighborhoods and areas
DUBAI_NEIGHBORHOODS = [
    "Downtown Dubai", "Palm Jumeirah", "Dubai Marina", "Jumeirah Beach Residence",
    "Business Bay", "Dubai Hills Estate", "Arabian Ranches", "Emirates Hills",
    "The Springs", "The Meadows", "Jumeirah Islands", "Jumeirah Park",
    "Dubai Silicon Oasis", "Dubai Sports City", "Dubai Production City",
    "Dubai Media City", "Dubai Internet City", "Dubai Knowledge Park",
    "Dubai Healthcare City", "Dubai International Financial Centre",
    "Al Barsha", "Al Quoz", "Al Safa", "Al Wasl", "Umm Suqeim",
    "Jumeirah", "Al Sufouh", "Al Manara", "Al Thanyah", "Al Hudaiba",
    "Al Satwa", "Al Karama", "Al Mankhool", "Al Raffa", "Al Jafiliya",
    "Al Qusais", "Al Nahda", "Al Qusais Industrial", "Al Warqa",
    "Mirdif", "Al Khawaneej", "Al Awir", "Al Lisaili", "Al Marmoom",
    "Al Faqa", "Al Ruwayyah", "Al Aweer", "Al Yalayis", "Al Ghadeer"
]

# Property types
PROPERTY_TYPES = [
    "Apartment", "Villa", "Townhouse", "Penthouse", "Studio",
    "Duplex", "Triplex", "Loft", "Garden Apartment", "Serviced Apartment"
]

# Developers
DEVELOPERS = [
    "Emaar Properties", "Nakheel", "Dubai Properties", "Meraas",
    "Damac Properties", "Sobha Realty", "Azizi Developments",
    "Deyaar Development", "Union Properties", "Al Habtoor Group",
    "Omniyat", "Select Group", "Ellington Properties", "Bloom Properties",
    "Aqua Properties", "Better Homes", "Chestertons", "Cluttons",
    "Knight Frank", "Savills"
]

# Amenities
AMENITIES = [
    "Swimming Pool", "Gym", "Spa", "Tennis Court", "Basketball Court",
    "Children's Playground", "BBQ Area", "Garden", "Balcony", "Terrace",
    "Parking", "Security", "Concierge", "24/7 Maintenance", "Pet Friendly",
    "Central AC", "Furnished", "Built-in Wardrobes", "Modern Kitchen",
    "Walk-in Closet", "Study Room", "Maid's Room", "Driver's Room"
]

# Transaction types
TRANSACTION_TYPES = ["Sale", "Rent", "Lease", "Sublease", "Assignment"]

# User roles
USER_ROLES = ["client", "agent", "employee", "admin", "manager", "supervisor"]

# Market data categories
MARKET_CATEGORIES = [
    "Residential", "Commercial", "Retail", "Industrial", "Hospitality",
    "Healthcare", "Education", "Mixed-use", "Land", "Off-plan"
]

def generate_properties_data(num_properties=500):
    """Generate comprehensive property data"""
    properties = []
    
    for i in range(num_properties):
        # Basic property info
        property_type = random.choice(PROPERTY_TYPES)
        neighborhood = random.choice(DUBAI_NEIGHBORHOODS)
        developer = random.choice(DEVELOPERS)
        
        # Price ranges based on property type and area
        if property_type in ["Villa", "Penthouse"]:
            base_price = random.randint(2000000, 15000000)
        elif property_type in ["Townhouse", "Duplex"]:
            base_price = random.randint(1500000, 8000000)
        else:
            base_price = random.randint(500000, 3000000)
        
        # Add some variation
        price = base_price + random.randint(-200000, 500000)
        
        # Size and rooms
        if property_type == "Studio":
            bedrooms = 0
            bathrooms = 1
            square_feet = random.randint(300, 600)
        elif property_type == "Apartment":
            bedrooms = random.randint(1, 4)
            bathrooms = random.randint(1, 3)
            square_feet = random.randint(600, 2000)
        else:
            bedrooms = random.randint(3, 7)
            bathrooms = random.randint(2, 6)
            square_feet = random.randint(2000, 8000)
        
        # Completion date
        completion_date = fake.date_between(start_date='-5y', end_date='+2y')
        
        # Views
        views = random.choice(["Sea View", "City View", "Garden View", "Pool View", "Street View", "Mountain View"])
        
        # Amenities (random selection)
        property_amenities = random.sample(AMENITIES, random.randint(3, 8))
        
        # Service charges
        service_charges = random.randint(5000, 50000)
        
        # Agent and agency
        agent = fake.name()
        agency = fake.company()
        
        # Address
        building_number = random.randint(1, 999)
        floor_number = random.randint(1, 50) if property_type != "Villa" else 1
        unit_number = random.randint(1, 999) if property_type != "Villa" else 1
        
        address = f"Building {building_number}, Floor {floor_number}, Unit {unit_number}, {neighborhood}, Dubai, UAE"
        
        property_data = {
            "id": i + 1,
            "address": address,
            "price_aed": price,
            "bedrooms": bedrooms,
            "bathrooms": bathrooms,
            "square_feet": square_feet,
            "property_type": property_type,
            "area": neighborhood,
            "developer": developer,
            "completion_date": completion_date.strftime("%Y-%m-%d"),
            "view": views,
            "amenities": ", ".join(property_amenities),
            "service_charges": service_charges,
            "agent": agent,
            "agency": agency,
            "status": random.choice(["Available", "Under Contract", "Sold", "Rented"]),
            "listing_date": fake.date_between(start_date='-1y', end_date='today').strftime("%Y-%m-%d"),
            "last_updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        properties.append(property_data)
    
    return properties

def generate_transactions_data(num_transactions=300):
    """Generate transaction history data"""
    transactions = []
    
    for i in range(num_transactions):
        transaction_type = random.choice(TRANSACTION_TYPES)
        property_id = random.randint(1, 500)
        
        # Price based on transaction type
        if transaction_type == "Sale":
            base_price = random.randint(800000, 12000000)
        else:
            base_price = random.randint(50000, 200000)  # Annual rent
        
        transaction_data = {
            "id": i + 1,
            "property_id": property_id,
            "transaction_type": transaction_type,
            "amount_aed": base_price,
            "transaction_date": fake.date_between(start_date='-2y', end_date='today').strftime("%Y-%m-%d"),
            "buyer_name": fake.name(),
            "seller_name": fake.name(),
            "agent_name": fake.name(),
            "agency": fake.company(),
            "commission_rate": random.uniform(2.0, 5.0),
            "commission_amount": base_price * random.uniform(0.02, 0.05),
            "status": random.choice(["Completed", "Pending", "Cancelled", "Under Review"]),
            "payment_method": random.choice(["Cash", "Bank Transfer", "Cheque", "Mortgage"]),
            "notes": fake.text(max_nb_chars=200)
        }
        
        transactions.append(transaction_data)
    
    return transactions

def generate_users_data(num_users=200):
    """Generate user data"""
    users = []
    
    for i in range(num_users):
        role = random.choice(USER_ROLES)
        
        user_data = {
            "id": i + 1,
            "username": fake.user_name(),
            "email": fake.email(),
            "first_name": fake.first_name(),
            "last_name": fake.last_name(),
            "phone": fake.phone_number(),
            "role": role,
            "department": random.choice(["Sales", "Marketing", "Operations", "Finance", "HR", "IT", "Legal", "Customer Service"]) if role in ["employee", "admin", "manager"] else None,
            "hire_date": fake.date_between(start_date='-3y', end_date='today').strftime("%Y-%m-%d"),
            "status": random.choice(["Active", "Inactive", "Suspended"]),
            "last_login": fake.date_time_between(start_date='-30d', end_date='now').strftime("%Y-%m-%d %H:%M:%S"),
            "created_at": fake.date_between(start_date='-1y', end_date='today').strftime("%Y-%m-%d")
        }
        
        users.append(user_data)
    
    return users

def generate_market_data(num_records=150):
    """Generate market analysis data"""
    market_data = []
    
    for i in range(num_records):
        category = random.choice(MARKET_CATEGORIES)
        neighborhood = random.choice(DUBAI_NEIGHBORHOODS)
        
        # Market metrics
        avg_price_per_sqft = random.randint(800, 3000)
        price_change_3m = random.uniform(-15.0, 25.0)
        price_change_6m = random.uniform(-20.0, 30.0)
        price_change_1y = random.uniform(-25.0, 40.0)
        
        market_record = {
            "id": i + 1,
            "category": category,
            "neighborhood": neighborhood,
            "avg_price_per_sqft": avg_price_per_sqft,
            "price_change_3m": round(price_change_3m, 2),
            "price_change_6m": round(price_change_6m, 2),
            "price_change_1y": round(price_change_1y, 2),
            "transaction_volume": random.randint(10, 500),
            "days_on_market": random.randint(15, 180),
            "supply_level": random.choice(["Low", "Medium", "High"]),
            "demand_level": random.choice(["Low", "Medium", "High"]),
            "report_date": fake.date_between(start_date='-6m', end_date='today').strftime("%Y-%m-%d"),
            "source": random.choice(["DLD", "RERA", "Property Monitor", "Chestertons", "Knight Frank"])
        }
        
        market_data.append(market_record)
    
    return market_data

def generate_employees_data(num_employees=100):
    """Generate employee data"""
    employees = []
    
    departments = ["Sales", "Marketing", "Operations", "Finance", "HR", "IT", "Legal", "Customer Service"]
    positions = ["Manager", "Senior", "Junior", "Assistant", "Director", "Coordinator", "Specialist"]
    
    for i in range(num_employees):
        department = random.choice(departments)
        position = random.choice(positions)
        
        employee_data = {
            "id": i + 1,
            "employee_id": f"EMP{str(i+1).zfill(4)}",
            "first_name": fake.first_name(),
            "last_name": fake.last_name(),
            "email": fake.email(),
            "phone": fake.phone_number(),
            "department": department,
            "position": position,
            "hire_date": fake.date_between(start_date='-5y', end_date='today').strftime("%Y-%m-%d"),
            "salary": random.randint(8000, 50000),
            "manager_id": random.randint(1, 20) if position != "Manager" else None,
            "status": random.choice(["Active", "On Leave", "Terminated"]),
            "performance_rating": random.uniform(1.0, 5.0),
            "specializations": random.sample(["Residential", "Commercial", "Luxury", "Off-plan", "Investment"], random.randint(1, 3))
        }
        
        employees.append(employee_data)
    
    return employees

def generate_legislations_data():
    """Generate legislation and policy data"""
    legislations = [
        {
            "id": 1,
            "title": "Dubai Real Estate Regulatory Authority (RERA) Regulations",
            "category": "Regulatory",
            "description": "Comprehensive regulations governing real estate transactions in Dubai",
            "effective_date": "2023-01-01",
            "status": "Active",
            "key_points": [
                "Property registration requirements",
                "Agent licensing standards",
                "Transaction documentation",
                "Dispute resolution procedures"
            ]
        },
        {
            "id": 2,
            "title": "Strata Law for Jointly Owned Properties",
            "category": "Property Management",
            "description": "Regulations for management of jointly owned properties and common areas",
            "effective_date": "2022-06-01",
            "status": "Active",
            "key_points": [
                "Owners association formation",
                "Service charge regulations",
                "Maintenance responsibilities",
                "Dispute resolution"
            ]
        },
        {
            "id": 3,
            "title": "Foreign Ownership Regulations",
            "category": "Investment",
            "description": "Rules governing foreign ownership of real estate in Dubai",
            "effective_date": "2021-03-01",
            "status": "Active",
            "key_points": [
                "Freehold areas designation",
                "Ownership restrictions",
                "Visa eligibility",
                "Tax implications"
            ]
        }
    ]
    
    return legislations

def create_csv_files():
    """Create CSV files with sample data"""
    print("Generating CSV files...")
    
    # Properties CSV
    properties = generate_properties_data(500)
    df_properties = pd.DataFrame(properties)
    df_properties.to_csv('data/properties.csv', index=False)
    print(f"Created properties.csv with {len(properties)} records")
    
    # Transactions CSV
    transactions = generate_transactions_data(300)
    df_transactions = pd.DataFrame(transactions)
    df_transactions.to_csv('data/transactions.csv', index=False)
    print(f"Created transactions.csv with {len(transactions)} records")
    
    # Users CSV
    users = generate_users_data(200)
    df_users = pd.DataFrame(users)
    df_users.to_csv('data/users.csv', index=False)
    print(f"Created users.csv with {len(users)} records")
    
    # Market Data CSV
    market_data = generate_market_data(150)
    df_market = pd.DataFrame(market_data)
    df_market.to_csv('data/market_data.csv', index=False)
    print(f"Created market_data.csv with {len(market_data)} records")
    
    # Employees CSV
    employees = generate_employees_data(100)
    df_employees = pd.DataFrame(employees)
    df_employees.to_csv('data/employees.csv', index=False)
    print(f"Created employees.csv with {len(employees)} records")

def create_excel_files():
    """Create Excel files with sample data"""
    print("Generating Excel files...")
    
    # Create workbook with multiple sheets
    wb = Workbook()
    
    # Properties sheet
    properties = generate_properties_data(200)
    ws_properties = wb.active
    ws_properties.title = "Properties"
    
    # Add headers
    headers = list(properties[0].keys())
    for col, header in enumerate(headers, 1):
        cell = ws_properties.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    
    # Add data
    for row, property_data in enumerate(properties, 2):
        for col, header in enumerate(headers, 1):
            ws_properties.cell(row=row, column=col, value=property_data[header])
    
    # Market Analysis sheet
    market_data = generate_market_data(100)
    ws_market = wb.create_sheet("Market Analysis")
    
    headers = list(market_data[0].keys())
    for col, header in enumerate(headers, 1):
        cell = ws_market.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
    
    for row, data in enumerate(market_data, 2):
        for col, header in enumerate(headers, 1):
            ws_market.cell(row=row, column=col, value=data[header])
    
    # Save workbook
    wb.save('data/real_estate_data.xlsx')
    print("Created real_estate_data.xlsx with multiple sheets")

def create_word_documents():
    """Create Word documents with sample data"""
    print("Generating Word documents...")
    
    # Company Policies Document
    doc = Document()
    doc.add_heading('Dubai Real Estate Company Policies', 0)
    
    policies = [
        ("Employee Code of Conduct", "All employees must maintain professional standards..."),
        ("Client Privacy Policy", "We are committed to protecting client information..."),
        ("Property Listing Standards", "All properties must meet quality standards..."),
        ("Commission Structure", "Standard commission rates and payment terms..."),
        ("Marketing Guidelines", "Brand guidelines and marketing standards...")
    ]
    
    for title, content in policies:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        doc.add_paragraph(f"Last updated: {fake.date_between(start_date='-6m', end_date='today').strftime('%Y-%m-%d')}")
        doc.add_paragraph("")
    
    doc.save('data/company_policies.docx')
    print("Created company_policies.docx")
    
    # Market Report Document
    doc = Document()
    doc.add_heading('Dubai Real Estate Market Report', 0)
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d')}")
    
    neighborhoods = random.sample(DUBAI_NEIGHBORHOODS, 10)
    for neighborhood in neighborhoods:
        doc.add_heading(f'{neighborhood} Market Analysis', level=1)
        doc.add_paragraph(f"Average price per sq ft: AED {random.randint(800, 3000)}")
        doc.add_paragraph(f"Price change (3 months): {random.uniform(-15, 25):.1f}%")
        doc.add_paragraph(f"Transaction volume: {random.randint(10, 500)} properties")
        doc.add_paragraph(f"Days on market: {random.randint(15, 180)} days")
        doc.add_paragraph("")
    
    doc.save('data/market_report.docx')
    print("Created market_report.docx")

def create_pdf_documents():
    """Create PDF documents with sample data"""
    print("Generating PDF documents...")
    
    # Property Brochure
    c = canvas.Canvas("data/property_brochure.pdf", pagesize=letter)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, 750, "Dubai Luxury Properties")
    c.setFont("Helvetica", 12)
    c.drawString(100, 720, f"Generated on: {datetime.now().strftime('%Y-%m-%d')}")
    
    y_position = 680
    for i in range(10):
        property_data = generate_properties_data(1)[0]
        c.drawString(100, y_position, f"Property {i+1}: {property_data['address']}")
        c.drawString(100, y_position-15, f"Price: AED {property_data['price_aed']:,}")
        c.drawString(100, y_position-30, f"Type: {property_data['property_type']} | Beds: {property_data['bedrooms']} | Baths: {property_data['bathrooms']}")
        y_position -= 60
        
        if y_position < 100:
            c.showPage()
            y_position = 750
    
    c.save()
    print("Created property_brochure.pdf")
    
    # Legal Document
    c = canvas.Canvas("data/legal_guidelines.pdf", pagesize=letter)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(100, 750, "Dubai Real Estate Legal Guidelines")
    
    c.setFont("Helvetica", 10)
    legal_text = [
        "1. Property Registration Requirements",
        "   - All properties must be registered with DLD",
        "   - Required documents: Title deed, NOC, etc.",
        "",
        "2. Agent Licensing",
        "   - All agents must be licensed by RERA",
        "   - Continuing education requirements apply",
        "",
        "3. Transaction Procedures",
        "   - Standard contracts must be used",
        "   - Escrow accounts for off-plan sales",
        "",
        "4. Dispute Resolution",
        "   - RERA dispute resolution center",
        "   - Arbitration procedures available"
    ]
    
    y_position = 700
    for line in legal_text:
        c.drawString(100, y_position, line)
        y_position -= 15
    
    c.save()
    print("Created legal_guidelines.pdf")

def create_json_files():
    """Create JSON files with sample data"""
    print("Generating JSON files...")
    
    # Neighborhood profiles
    neighborhoods = []
    for neighborhood in DUBAI_NEIGHBORHOODS[:20]:  # First 20 neighborhoods
        neighborhood_data = {
            "name": neighborhood,
            "average_price_per_sqft": random.randint(800, 3000),
            "popular_property_types": random.sample(PROPERTY_TYPES, 3),
            "top_developers": random.sample(DEVELOPERS, 5),
            "amenities": random.sample(AMENITIES, 8),
            "transportation": random.sample(["Metro", "Bus", "Tram", "Taxi", "Car"], 3),
            "schools": random.randint(2, 15),
            "hospitals": random.randint(1, 5),
            "shopping_centers": random.randint(1, 8),
            "parks": random.randint(1, 5),
            "description": f"{neighborhood} is a vibrant community in Dubai offering excellent amenities and connectivity."
        }
        neighborhoods.append(neighborhood_data)
    
    with open('data/neighborhoods.json', 'w') as f:
        json.dump(neighborhoods, f, indent=2)
    print("Created neighborhoods.json")
    
    # Company hierarchy
    hierarchy = {
        "company_name": "Dubai Real Estate Group",
        "structure": {
            "CEO": {
                "name": fake.name(),
                "department": "Executive",
                "direct_reports": ["COO", "CFO", "CTO"]
            },
            "COO": {
                "name": fake.name(),
                "department": "Operations",
                "direct_reports": ["Sales Director", "Marketing Director", "HR Director"]
            },
            "CFO": {
                "name": fake.name(),
                "department": "Finance",
                "direct_reports": ["Finance Manager", "Accounting Manager"]
            },
            "CTO": {
                "name": fake.name(),
                "department": "Technology",
                "direct_reports": ["IT Manager", "Development Manager"]
            }
        }
    }
    
    with open('data/company_hierarchy.json', 'w') as f:
        json.dump(hierarchy, f, indent=2)
    print("Created company_hierarchy.json")

def main():
    """Main function to generate all sample data"""
    print("Starting sample data generation...")
    
    # Create data directory if it doesn't exist
    os.makedirs('data', exist_ok=True)
    
    # Generate all data files
    create_csv_files()
    create_excel_files()
    create_word_documents()
    create_pdf_documents()
    create_json_files()
    
    print("\nSample data generation completed!")
    print("Generated files:")
    print("- CSV files: properties.csv, transactions.csv, users.csv, market_data.csv, employees.csv")
    print("- Excel file: real_estate_data.xlsx")
    print("- Word documents: company_policies.docx, market_report.docx")
    print("- PDF documents: property_brochure.pdf, legal_guidelines.pdf")
    print("- JSON files: neighborhoods.json, company_hierarchy.json")

if __name__ == "__main__":
    main()
